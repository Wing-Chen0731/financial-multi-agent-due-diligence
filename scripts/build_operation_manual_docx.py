from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "操作手册.md"
OUTPUT = ROOT / "金融多智能体系统_详细操作手册.docx"

NAVY = "16324F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F7F9FC"
GOLD = "A87127"
CODE_BG = "F4F6F8"
BORDER = "D7DEE7"
CJK_FONT = ".CJK Symbols Fallback SC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int], indent: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_run_font(run, name=CJK_FONT, size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    fonts.set(qn("w:cs"), CJK_FONT)
    fonts.set(qn("w:hint"), "eastAsia")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color=BORDER, size="6", space="4") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MUTED)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = CJK_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
    normal_fonts = normal._element.rPr.rFonts
    normal_fonts.set(qn("w:eastAsia"), CJK_FONT)
    normal_fonts.set(qn("w:cs"), CJK_FONT)
    normal_fonts.set(qn("w:hint"), "eastAsia")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = CJK_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
        style_fonts = style._element.rPr.rFonts
        style_fonts.set(qn("w:eastAsia"), CJK_FONT)
        style_fonts.set(qn("w:cs"), CJK_FONT)
        style_fonts.set(qn("w:hint"), "eastAsia")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = CJK_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
        style_fonts = style._element.rPr.rFonts
        style_fonts.set(qn("w:eastAsia"), CJK_FONT)
        style_fonts.set(qn("w:cs"), CJK_FONT)
        style_fonts.set(qn("w:hint"), "eastAsia")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = doc.styles.add_style("Manual Code", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_fonts = code._element.rPr.rFonts
    code_fonts.set(qn("w:eastAsia"), CJK_FONT)
    code_fonts.set(qn("w:cs"), CJK_FONT)
    code_fonts.set(qn("w:hint"), "eastAsia")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor.from_string("344054")
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(1)
    code.paragraph_format.space_after = Pt(1)
    code.paragraph_format.line_spacing = 1.08

    note = doc.styles.add_style("Manual Note", 1)
    note.font.name = CJK_FONT
    note._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
    note._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
    note_fonts = note._element.rPr.rFonts
    note_fonts.set(qn("w:eastAsia"), CJK_FONT)
    note_fonts.set(qn("w:cs"), CJK_FONT)
    note_fonts.set(qn("w:hint"), "eastAsia")
    note.font.size = Pt(10.5)
    note.font.color.rgb = RGBColor.from_string("475467")
    note.paragraph_format.left_indent = Inches(0.16)
    note.paragraph_format.right_indent = Inches(0.16)
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.line_spacing = 1.2

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("衡鉴 · 金融多智能体系统  |  操作手册")
    set_run_font(header_run, size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("内部作品集演示 · 第 ")
    set_run_font(footer_run, size=9, color=MUTED)
    add_page_field(footer)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("FINANCIAL MULTI-AGENT WORKBENCH")
    set_run_font(run, size=10, color=GOLD, bold=True)
    kicker.paragraph_format.space_after = Pt(22)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("衡鉴 · 金融多智能体系统")
    set_run_font(run, size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("详细操作手册")
    set_run_font(run, size=18, color=BLUE, bold=False)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(20)
    set_paragraph_border(rule, color=GOLD, size="10", space="1")

    metadata = doc.add_table(rows=4, cols=2)
    set_table_width(metadata, [2200, 7160], indent=120)
    metadata.style = "Table Grid"
    mark_header_row(metadata.rows[0])
    rows = [
        ("项目定位", "对公信贷尽调与合规辅助分析的多智能体作品集项目"),
        ("技术栈", "LangGraph · FastAPI · MCP JSON-RPC · Embedding RAG · Skills"),
        ("适用读者", "项目作者、面试官、同学、首次运行和验收人员"),
        ("版本", "v0.1.0 · 2026 年 8 月"),
    ]
    for row, (label, value) in zip(metadata.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_shading(row.cells[1], "FFFFFF")
        for idx, text in enumerate((label, value)):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            set_run_font(r, size=10.5, color=NAVY if idx == 0 else "475467", bold=idx == 0)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    note = doc.add_paragraph(style="Manual Note")
    set_paragraph_shading(note, CALLOUT)
    note.add_run("重要说明：").bold = True
    note.add_run("本项目使用本地演示 CRM、演示新闻和演示知识库，不代表真实银行生产环境。系统只提供辅助分析，不替代授信审批、合规判断或投资决策。")
    for run in note.runs:
        set_run_font(run, size=10.5, color="475467", bold=run.bold)

    doc.add_page_break()


def add_navigation(doc: Document, headings: list[str]) -> None:
    doc.add_heading("使用导航", level=1)
    p = doc.add_paragraph("这是一份面向实际试用和作品展示的操作手册。建议第一次运行时先阅读第 3 节 Demo 模式，再根据需要切换到本地 Ollama 或远程模型。Word 中可通过“导航窗格”按标题快速定位。")
    p.paragraph_format.space_after = Pt(10)
    doc.add_heading("快速路径", level=2)
    paths = [
        ("只看前端与流程", "启动 Demo 模式，不需要 API Key，也不需要下载模型。"),
        ("验证本地开源模型", "安装 Ollama，运行 setup 脚本，启动本地模型与向量 RAG。"),
        ("让同学不下载模型", "使用远程 OpenRouter / Hugging Face 模式，但每位使用者需要自己的 API Key。"),
        ("验证作品完整性", "依次检查 Web、API、MCP、RAG、Skill、测试和 Golden Dataset。"),
    ]
    for label, detail in paths:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{label}：")
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(detail)
        set_run_font(r, color=NAVY)

    doc.add_heading("章节导航", level=2)
    for heading in headings:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(heading)
        set_run_font(r, size=10.5, color="475467")
    doc.add_page_break()


def add_inline(paragraph, text: str) -> None:
    text = re.sub(r"<((?:https?://|mailto:)[^>]+)>", r"\1", text)
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, color=NAVY)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, color=NAVY, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color="344054")
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, color=NAVY, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, color=NAVY)


def add_code_block(doc: Document, code_lines: list[str], language: str) -> None:
    if language:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(language.upper())
        set_run_font(r, size=8, color=GOLD, bold=True)
    for line in code_lines:
        p = doc.add_paragraph(style="Manual Code")
        set_paragraph_shading(p, CODE_BG)
        add_inline(p, line if line else " ")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Manual Note")
    set_paragraph_shading(p, CALLOUT)
    add_inline(p, text)
    set_paragraph_border(p, color=LIGHT_BLUE, size="4", space="5")


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    if cols == 1:
        for row in rows:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, row[0])
        return
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    if cols == 2:
        widths = [2500, 6860]
    elif cols == 3:
        widths = [1600, 3300, 4460]
    else:
        widths = [9360 // cols for _ in range(cols)]
        widths[-1] += 9360 - sum(widths)
    set_table_width(table, widths, indent=120)
    for row_idx, row_data in enumerate(rows):
        for col_idx in range(cols):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            add_inline(p, row_data[col_idx] if col_idx < len(row_data) else "")
            for run in p.runs:
                set_run_font(run, size=9.5, color=NAVY, bold=row_idx == 0)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_table_row(line: str) -> list[str]:
    value = line.strip().strip("|")
    return [part.strip() for part in value.split("|")]


def render_markdown_body(doc: Document, source: str) -> list[str]:
    lines = source.replace("\r\n", "\n").split("\n")
    first_h2 = next((i for i, line in enumerate(lines) if line.startswith("## ")), 0)
    lines = lines[first_h2:]
    headings: list[str] = []
    index = 0
    in_code = False
    code_lines: list[str] = []
    code_lang = ""
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines, code_lang)
                in_code = False
                code_lines = []
                code_lang = ""
            else:
                in_code = True
                code_lang = line[3:].strip()
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1)) - 1
            title = re.sub(r"\s+#+$", "", heading.group(2)).strip()
            headings.append(title)
            doc.add_heading(title, level=level)
            index += 1
            continue
        if re.match(r"^\s*\|.*\|\s*$", line):
            table_lines = []
            while index < len(lines) and re.match(r"^\s*\|.*\|\s*$", lines[index]):
                candidate = parse_table_row(lines[index])
                if not all(re.fullmatch(r":?-+:?", cell) for cell in candidate):
                    table_lines.append(candidate)
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        if re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^\s*[-*]\s+", "", line))
            index += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\s*\d+\.\s+", "", line))
            index += 1
            continue
        if line.startswith(">"):
            add_note(doc, line[1:].strip())
            index += 1
            continue
        if re.match(r"^\s*---+\s*$", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(8)
            set_paragraph_border(p, color=LIGHT_BLUE, size="6", space="2")
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, line.rstrip())
        index += 1
    return headings


def add_appendix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("附录：未知客户资料缺失时的标准处理", level=1)
    add_note(doc, "这是本项目最重要的安全兜底流程：查不到客户不等于可以让模型猜测客户。系统必须先识别数据缺口，再要求补件或接入授权数据源。")
    doc.add_heading("标准状态机", level=2)
    steps = [
        ("1. 识别", "调用 crm.query_customer，确认客户是否存在于当前 CRM 数据源。"),
        ("2. 标记", "未找到时返回 status=not_found、label=资料待补充。"),
        ("3. 列清单", "返回主体资格、财务与现金流、债务与还款、担保与增信材料。"),
        ("4. 封锁结论", "can_form_lending_conclusion=false，禁止输出确定性的批准、拒绝或额度结论。"),
        ("5. 升级", "建议补录 CRM、上传材料或接入有授权的工商、征信、财务和核心系统。"),
        ("6. 复核", "材料齐备后由有权人员完成人工复核和最终判断。"),
    ]
    for label, detail in steps:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(f"{label}：")
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(detail)
        set_run_font(r, color=NAVY)

    doc.add_heading("测试输入与预期", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    set_table_width(table, [2700, 2300, 4360], indent=120)
    headers = ["测试输入", "状态", "预期行为"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        add_inline(p, text)
        for run in p.runs:
            set_run_font(run, size=9.5, color=NAVY, bold=True)
    test_rows = [
        ("海康威视科技有限公司", "not_found", "显示资料待补充卡片；置信度低；不形成授信结论"),
        ("示例科技有限公司", "found", "展示演示 CRM 记录；继续提示财务和担保材料需核验"),
        ("未提供公司名称", "not_found", "提示未识别客户，要求补充客户名称或资料"),
    ]
    for values in test_rows:
        row = table.add_row()
        for cell, text in zip(row.cells, values):
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, text)
            for run in p.runs:
                set_run_font(run, size=9.5, color=NAVY)
    set_table_width(table, [2700, 2300, 4360], indent=120)

    doc.add_heading("面试展示时的解释", level=2)
    add_note(doc, "可以这样说：系统不把“回答得像真的”当作成功标准。对于客户资料缺失的请求，系统优先保证事实边界、来源可追溯和人工可接管，因此会明确返回资料待补充，而不是使用模型记忆补写企业信息。")


def main() -> None:
    doc = Document()
    style_document(doc)
    doc.core_properties.title = "衡鉴 · 金融多智能体系统详细操作手册"
    doc.core_properties.subject = "LangGraph + FastAPI + MCP + RAG + Skills 操作与验收指南"
    doc.core_properties.author = "金融多智能体系统项目"
    doc.core_properties.keywords = "LangGraph, FastAPI, MCP, RAG, Ollama, Skill, 操作手册"

    add_cover(doc)
    source = SOURCE.read_text(encoding="utf-8")
    headings = [
        re.sub(r"^#{2,4}\s+", "", line).strip()
        for line in source.splitlines()
        if re.match(r"^#{2,4}\s+", line)
    ]
    add_navigation(doc, headings)
    render_markdown_body(doc, source)
    add_appendix(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
