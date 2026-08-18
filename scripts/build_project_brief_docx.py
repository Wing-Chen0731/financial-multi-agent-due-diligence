from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "金融多智能体系统_项目说明与系统架构说明书.docx"

NAVY = "16324F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
GOLD = "A87127"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F7F9FC"
GREEN = "EAF6EF"
GREEN_TEXT = "287A4D"
ORANGE = "FFF4E5"
ORANGE_TEXT = "8A5A00"
BORDER = "D7DEE7"
CODE_BG = "F4F6F8"
CJK_FONT = ".CJK Symbols Fallback SC"


def set_run_font(run, name=CJK_FONT, size=11, color=NAVY, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    fonts.set(qn("w:cs"), CJK_FONT)
    fonts.set(qn("w:hint"), "eastAsia")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int], indent: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_cell_text(cell, text: str, *, bold=False, color=NAVY, size=10.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(str(text))
    set_run_font(r, size=size, color=color, bold=bold)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], *, header_fill=LIGHT_BLUE, font_size=10.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, header_fill)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE, size=font_size)
    mark_header_row(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_data):
            set_cell_text(cell, value, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run, size=9, color=MUTED)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color=BORDER, size="6", space="4") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = CJK_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = CJK_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = CJK_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = doc.styles.add_style("Architecture Code", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    code._element.rPr.rFonts.set(qn("w:cs"), CJK_FONT)
    code._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    code.font.size = Pt(9.2)
    code.font.color.rgb = RGBColor.from_string("344054")
    code.paragraph_format.left_indent = Inches(0.16)
    code.paragraph_format.right_indent = Inches(0.16)
    code.paragraph_format.space_before = Pt(1)
    code.paragraph_format.space_after = Pt(1)
    code.paragraph_format.line_spacing = 1.08

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(2)
    run = header.add_run("衡鉴 · 金融多智能体系统  |  项目说明与系统架构")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("PTA作品集说明文档 · 第 ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(footer)


def add_para(doc, text: str, *, size=11, color=NAVY, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=6, before=0, line=1.25):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for run in p.runs:
        set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    set_keep_with_next(p)
    return p


def add_bullet(doc, text: str, *, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + 0.22 * level)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    r = p.add_run(text)
    set_run_font(r, size=11, color=NAVY)
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r, size=11, color=NAVY)
    return p


def start_numbered_list(doc: Document) -> int:
    """Create a fresh Word numbering instance so a new sequence starts at 1."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    level.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    level.append(jc)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_item(doc: Document, text: str, num_id: int):
    p = doc.add_paragraph(style="List Number")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)
    p_pr.append(num_pr)
    r = p.add_run(text)
    set_run_font(r, size=11, color=NAVY)
    return p


def add_callout(doc, label: str, text: str, *, fill=CALLOUT, label_color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(p, fill)
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.6, color=label_color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.6, color="475467")
    return p


def add_code(doc, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph(style="Architecture Code")
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.right_indent = Inches(0.16)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        set_paragraph_shading(p, CODE_BG)
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=9.2, color="344054")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_section_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(11)
    set_paragraph_border(p, color=LIGHT_BLUE, size="10", space="2")


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def cover(doc):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
    add_para(doc, "AI PRODUCT / SYSTEM ARCHITECTURE", size=10, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_para(doc, "某股份制银行对公信贷尽调智能助手", size=26, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.1)
    add_para(doc, "多智能体系统 · 项目说明与系统架构说明书", size=15, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_para(doc, "面向 PTA 项目作品集与面试答辩的系统化说明", size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1900, 7460])
    metadata = [
        ("项目时间", "2026.06 - 2026.08"),
        ("项目角色", "AI产品负责人（PTA项目）"),
        ("作品定位", "假想银行客户场景 · 公开信息与模拟推演 · 非生产系统"),
        ("技术主线", "LangGraph + MCP Tool Gateway + Skill Registry + Vector RAG + Web Workbench"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_text(row.cells[0], label, bold=True, color=DARK_BLUE, size=10.5)
        set_cell_text(row.cells[1], value, size=10.5)
    # Mark the first metadata row for Word accessibility tooling. The cover
    # table is a compact two-column key/value table, so the first row serves
    # as the repeated semantic anchor when read by assistive technology.
    mark_header_row(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    add_callout(doc, "阅读提示", "本说明书把“产品方案”“可运行 Demo”和“生产化差距”分开描述。文中“设计/规划”属于项目交付物，“当前实现”仅指本地作品可验证的能力，不等同于真实银行上线。", fill=ORANGE, label_color=ORANGE_TEXT)
    add_para(doc, "衡鉴 · Financial Multi-Agent Workbench", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=46, after=0)


def add_toc_like(doc):
    first = add_heading(doc, "如何阅读这份说明书", 1)
    first.paragraph_format.page_break_before = True
    add_para(doc, "这不是单纯的代码 README，而是一份把业务问题、产品判断、系统架构和可运行实现串起来的作品说明。面试展示时可先看第 1、3、4、9、12 章，再根据追问展开。", after=10)
    add_table(doc, ["阅读目标", "建议章节", "你要证明的能力"], [
        ["讲清楚为什么做", "第 1-2 章", "场景识别、价值判断、边界意识"],
        ["讲清楚怎么搭", "第 3-8 章", "多 Agent 架构、工具边界、RAG 与 Prompt"],
        ["讲清楚是否可信", "第 9-11 章", "金融安全、拒答、资料缺失、评测"],
        ["讲清楚如何落地", "第 12-14 章", "演示路径、生产差距、产品路线图"],
    ], [2200, 2400, 4760])
    add_heading(doc, "核心结论先看", 2)
    add_callout(doc, "一句话定位", "这是一个面向对公信贷客户经理的“决策辅助工作台”：用 LangGraph 编排事实采集、风险识别、合规依据检索和报告汇总，用 MCP/ToolGateway 约束外部工具边界，用 RAG 为制度类回答提供可追溯证据。", fill=GREEN, label_color=GREEN_TEXT)
    add_table(doc, ["层次", "本作品对应内容", "边界"], [
        ["产品层", "尽调流程、资料闸门、人工复核、前端工作台", "不替代客户经理或有权审批人"],
        ["Agent 层", "Supervisor + 4 个垂直 Agent", "输出分析与提示，不做审批决定"],
        ["工具层", "MCP JSON-RPC + ToolGateway 白名单", "当前为本地演示工具，生产需接授权系统"],
        ["知识层", "切片、Embedding、Chroma 持久化、混合检索", "当前知识库为演示材料，需正式文档治理"],
    ], [1800, 4700, 2860])
    page_break(doc)


def chapter_project(doc):
    add_heading(doc, "1. 项目背景与场景价值", 1)
    add_para(doc, "项目以某股份制银行对公信贷部为假想客户，围绕客户经理贷前尽调这一高频、资料密集、风险敏感的业务环节，完成一套多智能体协作系统的产品方案与可运行作品。项目时间为 2026.06 - 2026.08，项目角色为 AI 产品负责人（PTA 项目）。", after=7)
    add_para(doc, "需要特别说明：本项目基于公开信息、行业研究与模拟推演完成，CRM、行业新闻和知识库均为演示边界，不代表真实银行生产系统，也不构成授信审批意见。", color="475467", italic=True, after=10)
    add_heading(doc, "1.1 业务痛点与机会窗口", 2)
    add_table(doc, ["观察到的痛点", "对客户经理的影响", "产品机会"], [
        ["需手动查阅 CRM、企查查、内部制度等 5+ 系统", "信息分散、重复搬运、核验链路长", "以一个工作台聚合授权数据和制度证据"],
        ["单笔尽调耗时 8 小时以上", "大量时间耗在搜集、整理和格式化", "让 Agent 完成检索、摘要与结构化草稿"],
        ["信息遗漏导致约 12% 报告被退回", "返工、补件、人工复核压力上升", "用资料完整性检查和缺失材料清单前置拦截"],
        ["金融结论不可由模型随意推断", "幻觉会放大合规与信贷风险", "证据分层、红线 Prompt、人工复核状态"],
    ], [2600, 3100, 3660])
    add_heading(doc, "1.2 六维场景评估模型", 2)
    add_para(doc, "项目通过六维评估模型对 5 个候选场景进行横向比较，最终锁定对公信贷尽调。建议在面试中把“锁定场景”讲成可复用的决策框架，而不是凭直觉选题。", after=7)
    add_table(doc, ["维度", "判断问题", "尽调场景的典型表现"], [
        ["业务价值", "是否影响效率、质量或收入？", "直接影响尽调时长、报告退回和客户响应速度"],
        ["数据可得性", "是否有可授权、可追溯的数据？", "CRM、制度库、外部工商/征信均有接入路径"],
        ["流程标准化", "是否存在可编排的固定步骤？", "采集→风险→合规→报告的流程边界清晰"],
        ["模型适配度", "是否适合检索、归纳和辅助判断？", "适合证据整理与风险提示，不适合最终审批"],
        ["风险可控性", "能否定义红线、拒答和人工复核？", "可设置资料闸门、来源要求和高风险标记"],
        ["落地可行性", "是否能以 Demo 验证闭环？", "可用模拟 CRM、知识库和本地模型完成演示"],
    ], [1700, 3000, 4660])
    add_callout(doc, "价值假设", "方案目标是将单笔尽调耗时从 8 小时压缩至 40 分钟以内、将报告退回率从约 12% 降至 3% 以下。这是基于公开信息与模拟推演的预期测算，不应写成已上线实绩。", fill=ORANGE, label_color=ORANGE_TEXT)
    add_heading(doc, "1.3 产品目标与非目标", 2)
    add_table(doc, ["产品目标", "明确不做"], [
        ["减少跨系统查找和重复整理", "不替代真实身份认证、授权和数据治理"],
        ["输出结构化尽调草稿、风险提示和补件清单", "不自动批准/拒绝授信，不生成确定额度"],
        ["提供制度依据、来源和运行轨迹", "不把演示新闻当成正式事实核验"],
        ["将高风险任务标记为人工复核", "不提供个股买卖、投资组合或收益承诺"],
    ], [4680, 4680])


def chapter_positioning(doc):
    add_heading(doc, "2. 产品定位与用户旅程", 1)
    add_callout(doc, "产品定位", "衡鉴不是泛化聊天机器人，而是一个面向金融场景的决策辅助工作台。它把事实、证据、风险和决策权限拆开，让每一步都能被查看、解释和复核。", fill=GREEN, label_color=GREEN_TEXT)
    add_heading(doc, "2.1 目标用户与关键任务", 2)
    add_table(doc, ["角色", "主要任务", "系统提供的帮助"], [
        ["客户经理", "发起尽调、补件、撰写报告", "一站式查询、资料完整性提醒、报告草稿"],
        ["风险经理", "识别风险、复核证据", "风险提示、来源、缺失信息、工具轨迹"],
        ["合规人员", "核对制度与监管要求", "RAG 检索片段、文档来源、依据不足提示"],
        ["AI 产品/架构人员", "维护 Prompt、Skill、工具和评测", "版本信息、trace、Golden Dataset、运行状态"],
    ], [1800, 3000, 4560], font_size=9.6)
    add_heading(doc, "2.2 典型用户旅程：从发问到复核", 2)
    journey = [
        "客户经理输入：请对海康威视科技有限公司做一次授信尽调，重点看资料完整性、还款来源与担保安排。",
        "系统识别为 due_diligence，Supervisor 选择尽调相关 Skill，并保留 prompt_version 与 execution_mode。",
        "Data Collector 调用 CRM、知识库和行业新闻工具；只使用工具返回的事实，记录来源和审计轨迹。",
        "Risk Analyzer 输出风险点、缺失信息和后续核查项；不输出批准/拒绝结论。",
        "Compliance Checker 仅基于检索片段做辅助审查；没有依据时明确说明“未找到依据”。",
        "Report Writer 汇总成报告草稿；前端同时展示 trace、来源、置信度和人工复核标记。",
        "客户经理根据缺失材料清单补件，完成有权人员的人工复核与最终判断。",
    ]
    journey_num_id = start_numbered_list(doc)
    for item in journey:
        add_numbered_item(doc, item, journey_num_id)
    add_heading(doc, "2.3 输出物结构", 2)
    add_table(doc, ["输出区块", "内容", "证据等级"], [
        ["已知事实", "CRM 返回的客户资料、用户明确提供的财务数字", "可引用但仍需核验来源与时效"],
        ["检索依据", "制度库命中的片段、来源、页码/切片标识", "只代表检索到的辅助依据"],
        ["风险提示", "风险点、缺失项、核查建议", "分析性内容，不是授信结论"],
        ["下一步动作", "补件、接入授权数据源、人工复核", "流程建议，需业务人员执行"],
    ], [1800, 4200, 3360])


def chapter_architecture(doc):
    add_heading(doc, "3. 总体架构：1+1+N 的可解释协作系统", 1)
    add_para(doc, "“1+1+N”是本项目的核心架构表达：1 个智能中枢负责任务分流和流程控制，1 套 MCP/Tool Gateway 负责工具边界与审计，N 个垂直 Agent 负责数据采集、风险分析、合规审查和报告生成。", after=9)
    add_code(doc, [
        "用户 / 客户经理",
        "        │  Web Workbench / POST /chat",
        "        ▼",
        "┌──────────────────────────────────────────┐",
        "│  Supervisor · 任务分类 / Skill 选择 / 状态初始化 │",
        "└──────────────────────────────────────────┘",
        "        │ LangGraph FinancialState",
        "        ▼",
        "┌───────────────┐   ┌──────────────────────────┐",
        "│ Data Collector │──▶│ ToolGateway / MCP Gateway │",
        "└───────────────┘   └──────────────┬───────────┘",
        "        │                           │ 白名单 + 审计",
        "        ▼                           ▼",
        "┌───────────────┐   ┌──────────────────────────┐",
        "│ Risk Analyzer  │   │ CRM / Finance / News / RAG │",
        "└───────┬───────┘   └──────────────────────────┘",
        "        ▼",
        "┌──────────────────┐     ┌──────────────────────┐",
        "│ Compliance Checker│────▶│ Report Writer         │",
        "└──────────────────┘     └──────────┬───────────┘",
        "                                   ▼",
        "               报告 + 来源 + trace + 置信度 + 人工复核",
    ])
    add_heading(doc, "3.1 架构分层", 2)
    add_table(doc, ["层级", "组件", "职责", "当前作品证据"], [
        ["交互层", "Web Workbench", "输入问题、显示报告、trace、来源与运行状态", "web/index.html + web/app.js"],
        ["编排层", "LangGraph StateGraph", "定义节点、边和 FinancialState 的流转", "src/graph.py"],
        ["中枢层", "Supervisor", "任务类型识别、Skill 推荐、运行时信息", "SUPERVISOR_PROMPT + 规则兜底"],
        ["Agent 层", "4 个垂直 Agent", "采集、风险、合规、报告", "5 节点顺序链路可追踪"],
        ["工具层", "ToolGateway / MCP", "工具白名单、Schema、调用和审计", "src/mcp/gateway.py + /mcp"],
        ["知识层", "Loader / Embedding / Chroma / Retriever", "切片、向量化、持久化、混合检索", "src/rag/*"],
        ["治理层", "Prompt、Skill、评测、人工复核", "限制模型权限、验证输出、形成闭环", "src/skills + evals + state 字段"],
    ], [1300, 2200, 3300, 2560], font_size=9.6)
    add_heading(doc, "3.2 架构决策与取舍", 2)
    add_bullet(doc, "采用顺序图而非自由调用：Demo 更容易观测，后续可以再引入条件分支、并行节点和重试策略。")
    add_bullet(doc, "采用工具白名单而非让模型直接访问数据：模型只能提出工具调用，数据访问权由 Gateway 决定。")
    add_bullet(doc, "采用“模型能力 + 确定性安全规则”：高风险意图、未知客户和投资请求的安全规则具有更高优先级。")
    add_bullet(doc, "采用“向量检索 + 词法融合 + 显式 fallback”：保证本地演示可用，同时不把词法降级冒充成向量 RAG。")


def chapter_langgraph(doc):
    add_heading(doc, "4. LangGraph 工作流与共享状态", 1)
    add_para(doc, "LangGraph 在本项目中不是一个装饰性的技术名词，而是把 Agent 协作协议落成可执行状态图的核心编排层。每个节点只负责一类工作，并通过 FinancialState 将结构化结果交给下一个节点。", after=8)
    add_heading(doc, "4.1 节点与边", 2)
    add_table(doc, ["节点", "输入", "输出", "失败/边界处理"], [
        ["supervisor", "user_query", "task_type、selected_skills、prompt_version", "模型 JSON 解析失败时回到规则分类"],
        ["data_collector", "任务类型、客户名、Skill context", "CRM、新闻、RAG、财务指标、tool_trace", "客户未找到时进入资料待补充闸门"],
        ["risk_analyzer", "采集事实与行业新闻", "risk_analysis", "fast 模式使用可审计规则；full 模式模型失败有安全文本"],
        ["compliance_checker", "retrieved_context", "compliance_review", "没有依据时输出 no_evidence"],
        ["report_writer", "全部中间结果", "final_report、confidence、need_human_review", "未知客户优先输出补件报告，禁止确定性授信结论"],
    ], [1700, 2700, 2800, 2160], font_size=9.6)
    add_code(doc, [
        "START → supervisor → data_collector → risk_analyzer",
        "      → compliance_checker → report_writer → END",
        "",
        "FinancialState = {",
        "  user_query, task_type, selected_skills, skill_context,",
        "  collected_data, retrieved_context, risk_analysis,",
        "  compliance_review, final_report, confidence,",
        "  need_human_review, trace, tool_trace, runtime",
        "}",
    ])
    add_heading(doc, "4.2 关键状态字段", 2)
    add_table(doc, ["字段", "产品含义", "前端/面试价值"], [
        ["task_type", "general_chat / due_diligence / compliance_query / investment_query", "解释为什么走这条工作流"],
        ["selected_skills", "本次请求被推荐的 Skill", "证明 Skill 不是静态文档，而是进入运行上下文"],
        ["customer_data_status", "found / not_found + 缺失材料 + 下一步", "把资料完整性做成机器可读闸门"],
        ["retrieved_context", "知识库片段、来源、score、retrieval_mode", "实现引用展示和 RAG 透明度"],
        ["tool_trace / trace", "工具调用轨迹与节点轨迹", "用于调试、审计和面试演示"],
        ["confidence / need_human_review", "置信度和人工复核标记", "把结果不确定性显式呈现"],
    ], [2300, 4200, 2860], font_size=9.8)
    add_callout(doc, "核心设计", "报告生成节点不是“最后让大模型自由发挥”，而是先检查任务类型、客户资料状态、证据和高风险标记，再决定输出格式。未知客户路径会优先级高于普通报告生成。", fill=ORANGE, label_color=ORANGE_TEXT)


def chapter_agents(doc):
    add_heading(doc, "5. Agent 角色、Prompt 与协作协议", 1)
    add_para(doc, "项目为 4 个垂直 Agent 编写了带角色设定、工作流程、红线规则和输出格式的 System Prompt，并保留 Prompt 版本号。当前代码中的 PROMPT_VERSION 为 v1.1.0。", after=8)
    add_heading(doc, "5.1 Agent 角色矩阵", 2)
    add_table(doc, ["Agent", "负责什么", "不负责什么", "关键输出"], [
        ["Supervisor", "识别任务类型、选择 Skill、初始化运行状态", "不做审批、不做投资决定", "task_type、selected_skills"],
        ["Data Collector", "整理工具返回的 CRM、新闻、RAG 事实", "不补写不存在的数据", "collected_data、sources、tool_trace"],
        ["Risk Analyzer", "基于事实识别风险点、缺失信息、核查项", "不输出批准/拒绝授信结论", "risk_analysis"],
        ["Compliance Checker", "只对照检索片段做辅助审查", "不把有限片段说成最终合规结论", "compliance_review"],
        ["Report Writer", "汇总已知事实、依据、风险和下一步动作", "不生成额度、批准/拒绝或投资建议", "final_report、confidence"],
    ], [1700, 2700, 2700, 2260], font_size=9.7)
    add_heading(doc, "5.2 Prompt 的共同约束", 2)
    for item in [
        "事实优先：只整理工具和用户明确提供的数据，标记来源、时间和演示数据限制。",
        "证据不足即说明不足：没有知识库依据时不能用模型记忆补齐制度内容。",
        "结论降级：将“批准/拒绝”降级为“风险提示/核查建议/人工复核”。",
        "输出结构化：区分已知事实、检索依据、风险提示、缺失材料和下一步动作。",
        "安全规则优先：投资推荐、未知客户和高风险事项不由模型自由决定。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.3 适合面试展示的 Prompt 迭代逻辑", 2)
    add_table(doc, ["版本阶段", "主要问题", "迭代方向"], [
        ["v1.0", "能回答，但输出容易混合事实与推断", "加入事实/依据/风险/动作四段式"],
        ["v1.1", "金融场景边界不够强，未知客户可能被误解", "增加禁止编造、资料闸门和人工复核要求"],
        ["下一阶段", "不同模型输出稳定性和引用一致性需验证", "增加结构化 JSON、schema 校验、离线评测和重试"],
    ], [1600, 3500, 4260])


def chapter_mcp(doc):
    add_heading(doc, "6. MCP 与工具网关：让 Agent 有边界地行动", 1)
    add_para(doc, "MCP 在本项目中承担“工具协议和接入边界”的角色。ToolGateway 维护白名单、工具 Schema 和审计轨迹；/mcp 接口提供 initialize、tools/list、tools/call 的 JSON-RPC 演示入口。", after=8)
    add_heading(doc, "6.1 工具目录", 2)
    add_table(doc, ["工具名", "业务用途", "数据边界", "生产化接入方向"], [
        ["crm.query_customer", "查询演示 CRM 客户资料", "当前为本地模拟客户数据，不推断未返回事实", "接入 CRM API / OAuth / RBAC"],
        ["finance.calculate_ratios", "按用户输入计算利润率、债务收入比", "只计算，不验证期间、币种、审计状态", "接入财务数据服务并增加口径校验"],
        ["industry.search_news", "查询演示行业新闻线索", "不替代正式新闻、舆情或事实核验", "接入授权新闻源和时间有效性校验"],
        ["knowledge.retrieve", "检索制度与金融知识库", "返回带来源的辅助依据，不等于最终合规意见", "增加权限过滤、rerank、文档版本治理"],
    ], [2200, 2600, 2600, 1960], font_size=9.5)
    add_heading(doc, "6.2 MCP 调用边界", 2)
    add_code(doc, [
        "POST /mcp",
        "{\"jsonrpc\":\"2.0\",\"id\":1,\"method\":\"tools/list\",\"params\":{}}",
        "",
        "模型/Agent → 请求工具 → ToolGateway 白名单 → 具体工具 → 结果 + audit_log",
        "                                      └── 未授权工具：PermissionError",
    ])
    add_table(doc, ["控制点", "当前实现", "还需补齐"], [
        ["白名单", "仅允许 4 个工具名", "按租户/角色/数据域做动态权限"],
        ["Schema", "tools/list 返回 inputSchema", "增加 schema 校验、版本和字段脱敏声明"],
        ["审计", "ToolGateway 记录 tool、params、ok", "持久化、不可篡改、关联 user/session/request id"],
        ["认证", "可选 Bearer Token（MCP_AUTH_TOKEN）", "企业 SSO、OAuth、密钥轮换和限流"],
        ["容错", "工具异常返回用户安全错误", "超时、重试、熔断、降级与告警"],
    ], [1700, 3700, 3960], font_size=9.8)


def chapter_skill(doc):
    add_heading(doc, "7. Skill 系统：把业务规则做成版本化能力包", 1)
    add_para(doc, "Skill 不是散落在 Prompt 里的几句提示，而是位于 src/skills/**/SKILL.md 的 Markdown 能力规范。SkillRegistry 会读取 frontmatter，按触发词推荐 Skill，再把 Skill 内容注入 Agent 上下文。", after=8)
    add_heading(doc, "7.1 当前 Skill 与简历描述的映射", 2)
    add_table(doc, ["作品中的 Skill", "对应业务能力", "简历可表达的能力", "风险等级"], [
        ["customer_due_diligence", "企业授信尽调", "CRM 查询 / 尽调流程 Skill", "high"],
        ["compliance_knowledge_review", "制度合规审查", "RAG 知识库检索 Skill", "high"],
        ["industry_news_search", "行业新闻检索", "行业新闻检索 Skill", "medium"],
        ["financial_ratio_analysis", "基础财务指标", "财务指标分析 Skill", "medium"],
        ["general_finance_guardrail", "金融安全护栏", "高风险拒答与人工复核", "high"],
    ], [2500, 2600, 2800, 1460], font_size=9.7)
    add_heading(doc, "7.2 Skill 的标准结构", 2)
    add_code(doc, [
        "---",
        "name: customer_due_diligence",
        "description: 企业授信尽调流程与事实边界",
        "risk_level: high",
        "triggers: 尽调, 授信, 贷款, 贷前",
        "tools: crm.query_customer, knowledge.retrieve",
        "---",
        "目标 / 触发条件 / 输入输出 / 执行步骤 / 禁止事项 / 失败处理 / 示例",
    ])
    add_heading(doc, "7.3 Skill 的产品价值", 2)
    add_bullet(doc, "复用：同类金融任务可以共享规则，不必为每个 Agent 重写整套 Prompt。")
    add_bullet(doc, "可维护：业务规则以 Markdown 版本化，产品、合规和研发可以协同审阅。")
    add_bullet(doc, "可观测：selected_skills 会进入 API 返回和前端 trace，面试时能证明它真的参与运行。")
    add_bullet(doc, "可评测：可以按 Skill 组织 Golden Dataset，分开观察尽调、合规、财务和安全护栏效果。")


def chapter_rag(doc):
    add_heading(doc, "8. 完整向量 RAG：从文档到可引用依据", 1)
    add_para(doc, "本项目的 RAG 不是简单的本地词法检索。完整向量路径包括：文档加载 → 结构化切片 → Embedding → Chroma 持久化 → 相似度检索 → 与词法结果融合 → 返回来源和片段。", after=8)
    add_code(doc, [
        "data/knowledge_base/*.txt, *.md, *.pdf",
        "        │ loader：归一化文本 + chunk_size 800 + overlap 120",
        "        ▼",
        "Embedding：Ollama qwen3-embedding:0.6b / BGE / OpenAI-compatible remote",
        "        │",
        "        ▼",
        "Chroma PersistentClient → data/vector_db → financial_knowledge",
        "        │",
        "        ▼",
        "query embedding → cosine similarity → top_k chunks",
        "        │                         ↘ lexical overlap（auto 模式）",
        "        ▼",
        "hybrid score = 0.7 × vector_score + 0.3 × normalized_lexical_score",
        "        │",
        "        ▼",
        "snippet + source + page + chunk + score + retrieval_mode",
    ])
    add_heading(doc, "8.1 RAG 组件说明", 2)
    add_table(doc, ["组件", "当前实现", "作品表达"], [
        ["切片", "支持 txt/md/PDF；默认 chunk_size=800、overlap=120，并优先在段落/句号/分号处切分", "不是把整份制度文档一次塞进上下文"],
        ["Embedding", "Ollama、本地 SentenceTransformer、OpenAI-compatible 远程三类适配器", "Embedding 与聊天模型解耦"],
        ["向量库", "Chroma PersistentClient，默认 data/vector_db", "具备持久化索引，不是每次临时计算"],
        ["检索", "vector / lexical / auto 三种模式；auto 做混合融合", "明确区分完整向量能力和降级能力"],
        ["引用", "source、page、chunk、score、retrieval_mode", "让合规 Agent 和前端可以展示证据链"],
    ], [1700, 4800, 3160], font_size=9.6)
    add_heading(doc, "8.2 RAG 模式与分享策略", 2)
    add_table(doc, ["模式", "需要什么", "适合场景", "必须如何表述"], [
        ["vector", "Embedding 服务 + Chroma 索引", "验证完整向量 RAG", "明确说是向量检索"],
        ["auto", "优先向量，失败后 lexical fallback", "日常本地演示", "前端展示实际 retrieval_mode"],
        ["lexical", "仅知识库文本和词法打分", "同学零模型分享、无 Embedding 配额", "只能说是词法检索，不得宣传成向量 RAG"],
    ], [1500, 2700, 2460, 2700], font_size=9.5)
    add_callout(doc, "数据治理提醒", "当前 data/knowledge_base 中的材料是演示样例；用户简历中“50+内部制度文件”属于方案设计目标。正式版本还需要文档版本、生效日期、适用机构、权限标签、失效归档和增量索引。", fill=ORANGE, label_color=ORANGE_TEXT)
    add_heading(doc, "8.3 RAG 质量指标", 2)
    add_bullet(doc, "召回：Top-k 命中率、Recall@k、制度条款覆盖率。")
    add_bullet(doc, "排序：MRR、nDCG、rerank 前后对比。")
    add_bullet(doc, "引用：引用准确率、引用完整性、来源可追溯率。")
    add_bullet(doc, "安全：无依据拒答率、过期制度拦截率、越权文档召回率。")


def chapter_safety(doc):
    add_heading(doc, "9. 金融安全边界与异常兜底", 1)
    add_para(doc, "金融场景的核心不是让模型“更敢回答”，而是让系统在证据不足、客户不存在或问题高风险时知道何时停下来。作品把这些状态做成可观察的产品能力。", after=8)
    add_heading(doc, "9.1 未知客户的标准处理", 2)
    add_code(doc, [
        "客户经理输入客户名",
        "        ↓",
        "CRM 查询 → found = false",
        "        ↓",
        "customer_data_status = not_found / 资料待补充",
        "        ↓",
        "列出主体资格、财务现金流、债务还款、担保增信材料",
        "        ↓",
        "不编造企业信息 · 不输出确定性授信结论 · 标记人工复核",
        "        ↓",
        "补录 CRM / 上传材料 / 接入授权数据源 → 重新核验",
    ])
    add_table(doc, ["状态", "系统允许输出", "系统禁止输出"], [
        ["CRM found=false", "资料待补充、缺失材料、下一步动作、低置信度", "企业注册资本、行业地位、流水、债务等模型记忆事实"],
        ["RAG 无命中", "未找到可引用依据、建议补充制度材料", "把模型常识冒充制度条款"],
        ["投资类请求", "信息核查框架和人工复核提醒", "个股买卖、投资组合、收益承诺"],
        ["财务数字来自用户输入", "透明公式和口径限制", "把计算结果当成已核验财务结论"],
    ], [1800, 3600, 3960], font_size=9.7)
    add_heading(doc, "9.2 风险控制矩阵", 2)
    add_table(doc, ["风险", "控制手段", "可观察证据"], [
        ["幻觉/编造", "Prompt 禁止补写 + CRM found 闸门 + 事实/推断分层", "customer_data_status、报告中的边界说明"],
        ["依据不足", "RAG 无命中即 no_evidence + 合规人工复核", "compliance_review.status"],
        ["工具越权", "ToolGateway 白名单 + MCP Schema", "tools/list、PermissionError、tool_trace"],
        ["结论越界", "高风险任务统一 need_human_review=true", "confidence、need_human_review"],
        ["服务不可用", "demo/fast/remote/local 多档位 + 显式降级", "readiness、runtime、retrieval_mode"],
    ], [1800, 3900, 3660], font_size=9.6)
    add_callout(doc, "面试表述", "我没有把模型能力当成最终决策能力，而是通过数据闸门、证据检索、工具白名单和人工复核把可用范围收窄。这是金融 AI 产品设计中比“回答很像人”更重要的可信度。", fill=GREEN, label_color=GREEN_TEXT)


def chapter_ui(doc):
    add_heading(doc, "10. 前端工作台与可解释体验", 1)
    add_para(doc, "前端不是单纯的聊天框，而是把系统运行状态呈现给业务人员和面试官的“可解释入口”。打开首页后，用户可以选择企业授信尽调、合规问答和安全边界测试等场景，并查看结果旁的过程信息。", after=8)
    add_heading(doc, "10.1 页面信息架构", 2)
    add_table(doc, ["区域", "展示内容", "产品价值"], [
        ["模型与运行时", "provider、model、execution_mode、readiness", "知道当前用的是 demo、本地还是远程模型"],
        ["工作区/场景", "预置问题、输入框、运行按钮", "让面试演示有明确路径"],
        ["结果区", "结构化报告、资料状态、风险提示、下一步", "让输出可执行而不是一段散文"],
        ["Trace 区", "Supervisor、采集、风险、合规、报告节点", "证明多 Agent 真的执行过"],
        ["证据区", "来源、score、retrieval_mode、工具轨迹", "支撑可追溯与人工核验"],
    ], [1800, 3900, 3660], font_size=9.7)
    add_heading(doc, "10.2 三个推荐演示场景", 2)
    add_table(doc, ["场景", "示例输入", "重点展示"], [
        ["企业授信尽调", "请对海康威视科技有限公司做一次授信尽调，重点看资料完整性、还款来源与担保安排。", "5 节点 trace、CRM、RAG、风险提示、人工复核"],
        ["未知客户资料闸门", "请对星海未来科技有限公司做授信尽调。", "CRM 未找到→资料待补充→缺失材料→不编造"],
        ["合规与安全边界", "请给我推荐一只股票，保证三个月收益。", "Supervisor 分类、拒答、风险边界与人工复核"],
    ], [1900, 4800, 2660], font_size=9.4)
    add_heading(doc, "10.3 体验上的下一步", 2)
    add_bullet(doc, "将资料缺失状态从文本提示升级为可勾选补件清单，并支持上传文件。")
    add_bullet(doc, "将 trace 从节点列表升级为时间线，展示每一步耗时、输入摘要和输出状态。")
    add_bullet(doc, "在引用片段旁显示制度版本、生效日期和命中理由。")
    add_bullet(doc, "增加“提交人工复核/生成补件任务”按钮，形成从分析到流程的闭环。")


def chapter_eval(doc):
    add_heading(doc, "11. 评测体系与 Golden Dataset", 1)
    add_para(doc, "项目设计了 50 个测试用例，覆盖正常、边界和红线三类场景。Golden Dataset 的作用不是为了证明模型每次都能答对，而是验证它是否在高风险情况下按照产品规则拒答、降级或要求人工复核。", after=8)
    add_heading(doc, "11.1 测试用例分层", 2)
    add_table(doc, ["类别", "建议覆盖", "验收重点"], [
        ["正常场景", "企业尽调、制度问答、财务指标、行业新闻", "分类正确、工具调用正确、报告结构完整"],
        ["边界场景", "客户名缺失、CRM 未找到、RAG 无命中、数据不完整", "不编造、列出缺失材料、显式降级"],
        ["红线场景", "投资推荐、收益承诺、授信批准/拒绝、越权工具", "拒答或转人工，不能输出越界结论"],
    ], [1800, 4300, 3260])
    add_heading(doc, "11.2 建议评测指标", 2)
    add_table(doc, ["指标", "定义", "目标/解释"], [
        ["任务分类准确率", "task_type 是否与人工标注一致", "评估 Supervisor 的分流可靠性"],
        ["工具选择准确率", "是否调用正确工具、是否避免多余工具", "评估 MCP/Skill 协作边界"],
        ["事实一致性", "输出事实是否能在工具/用户输入中找到", "未知客户用例必须为 100% 不编造"],
        ["引用准确率", "引用片段是否支持回答中的制度表述", "评估 RAG 质量，不只看召回数量"],
        ["安全拒答率", "红线请求是否拒答/转人工", "安全优先于回答完整度"],
        ["人工复核召回率", "高风险/低置信请求是否正确标记", "保证系统不把风险藏在文本里"],
    ], [2200, 4100, 3060], font_size=9.7)
    add_heading(doc, "11.3 面试时如何展示评测思路", 2)
    add_para(doc, "可以选 3 个样例：一个正常尽调、一个未知客户、一个投资推荐。先展示系统结果，再说明 Golden Dataset 如何把这些行为固化成回归用例，最后指出正式生产还需要基于真实脱敏数据重新标注和评测。", after=7)


def chapter_deploy(doc):
    add_heading(doc, "12. 当前实现、启动方式与交付边界", 1)
    add_para(doc, "本章用于把“架构说明”落到作品可以实际运行的目录和命令上。它同时帮助面试官理解：哪些能力已经在本地 Demo 中可验证，哪些仍属于生产化规划。", after=8)
    add_heading(doc, "12.1 代码目录与架构映射", 2)
    add_table(doc, ["目录/文件", "对应模块", "说明"], [
        ["src/graph.py", "LangGraph 编排", "5 个节点、FinancialState、路由与报告生成"],
        ["src/agents/prompts.py", "System Prompt", "Supervisor、采集、风险、合规、报告 Prompt"],
        ["src/mcp/gateway.py", "MCP/工具边界", "4 个工具、Schema、白名单与审计"],
        ["src/skills/**/SKILL.md", "Skill 系统", "5 个可加载的金融 Skill"],
        ["src/rag/", "完整 RAG", "loader、embeddings、vector_store、retriever、service"],
        ["src/api/server.py", "服务入口", "Web、/health、/readiness、/chat、/mcp"],
        ["web/", "交互前端", "工作台页面和展示逻辑"],
        ["evals/ + tests/", "评测与回归", "Golden Dataset、RAG/Graph/Skill 测试"],
    ], [2500, 2600, 4260], font_size=9.8)
    add_heading(doc, "12.2 三种运行档位", 2)
    add_table(doc, ["档位", "启动方式", "模型/依赖", "适合展示"], [
        ["Demo", "./scripts/start.sh --demo", "无需 API Key 或模型", "先看前端、trace、边界和作品结构"],
        ["本地开源模型", "./scripts/setup.sh → ./scripts/start.sh", "Ollama + qwen3:1.7b / Embedding", "展示开源模型与完整本地链路"],
        ["远程模型", "./scripts/start_remote.sh", "OpenRouter/Hugging Face 等 Key", "同学不下载本地大模型，直接体验"],
    ], [2100, 3000, 2600, 1660], font_size=9.5)
    add_heading(doc, "12.3 API 入口", 2)
    add_code(doc, [
        "GET  /health       进程存活与运行时信息",
        "GET  /readiness    模型 / RAG 就绪状态",
        "POST /chat         运行 LangGraph，返回报告、trace、来源和状态",
        "POST /mcp          JSON-RPC initialize / tools/list / tools/call",
        "",
        "POST /chat body: {\"message\":\"请对星海未来科技有限公司做授信尽调\"}",
    ])
    add_heading(doc, "12.4 从 PTA 作品到生产系统的差距", 2)
    add_table(doc, ["当前作品已具备", "生产化还需补齐"], [
        ["完整 5 节点 LangGraph 链路和前端 trace", "持久化会话、任务重试、超时、熔断、分布式执行"],
        ["本地演示 CRM、新闻、知识库工具", "真实系统 OAuth/RBAC、字段级脱敏、授权记录和审计"],
        ["完整向量 RAG 与 lexical fallback", "制度文档治理、增量索引、权限过滤、rerank、引用评测"],
        ["Prompt/Skill 版本化和 50 条 Golden Dataset", "脱敏真实数据、线上监控、模型漂移和回归门禁"],
        ["人工复核状态标记", "工单、审批、补件和结果回流闭环"],
        ["本地/远程/无模型三种分享方式", "Docker、CI/CD、企业模型网关、成本和容量治理"],
    ], [4680, 4680], font_size=9.6)
    add_callout(doc, "交付边界", "作品的可信表达不是“已经上线银行生产环境”，而是“完成了可运行的方案验证，并清楚定义了从 Demo 到生产的补齐项”。这会让项目经历更真实，也更经得起面试追问。", fill=ORANGE, label_color=ORANGE_TEXT)


def chapter_interview(doc):
    add_heading(doc, "13. 面试答辩叙事与项目成果", 1)
    add_heading(doc, "13.1 项目成果的完整表达", 2)
    script_num_id = start_numbered_list(doc)
    for item in [
        "围绕对公信贷尽调完成场景评估与 ROI 测算，形成《场景评估与 ROI 测算报告》。",
        "设计“1 个智能中枢 + 1 套 MCP 网关 + N 个垂直 Agent”的协作架构，明确职责边界、输入输出和协作协议。",
        "为 4 个核心 Agent 编写 System Prompt，形成角色设定、工作流程、红线规则和输出格式，并保留版本信息。",
        "设计并实现 5 个可加载金融 Skill，覆盖企业尽调、合规审查、行业新闻、财务指标和金融安全护栏。",
        "设计完整向量 RAG 方案，包含层级切片、Embedding、Chroma 持久化、混合检索、来源展示和 fallback。",
        "产出 AI-PRD、Golden Dataset、项目答辩材料、操作手册和本次系统说明书，形成可展示作品闭环。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "13.2 三分钟答辩脚本", 2)
    for item in [
        "背景：客户经理要跨 5+ 系统完成尽调，单笔 8 小时以上，信息遗漏造成约 12% 报告退回。",
        "判断：尽调同时具备高价值、流程标准化、数据可接入和风险可控四个条件，适合作为 AI 辅助场景。",
        "方案：采用 1+1+N 架构，LangGraph 编排 Supervisor、采集、风险、合规、报告五个节点；MCP/ToolGateway 控制工具访问。",
        "关键设计：不让模型直接下授信结论，所有输出分成事实、依据、风险、缺失材料和下一步，并设置人工复核。",
        "演示：先跑正常尽调，再跑 CRM 未找到的客户，最后跑投资推荐红线，展示系统如何“知道何时不能回答”。",
        "边界：当前是公开信息和模拟推演的 PTA 作品，生产化还要补充真实授权、文档治理、评测和审批闭环。",
    ]:
        add_numbered_item(doc, item, script_num_id)
    add_heading(doc, "13.3 典型追问与回答方向", 2)
    add_table(doc, ["面试官追问", "回答重点"], [
        ["为什么不用一个大模型直接回答？", "金融场景需要可追溯、可审计和可复核；多 Agent 是职责拆分，不是为了堆模型。"],
        ["Supervisor Prompt 是否真的使用？", "full 模式调用 Supervisor Prompt，fast 模式用确定性规则保证速度；两者都有 trace，安全规则优先。"],
        ["RAG 是不是只有关键词检索？", "完整路径有切片、Embedding、Chroma 持久化和向量查询；auto 可融合 lexical，lexical 仅作为显式降级。"],
        ["客户不存在时模型会不会编造？", "CRM found=false 触发资料待补充状态，固定列缺失材料，不生成确定性授信结论。"],
        ["为什么项目不是生产系统？", "这是 PTA 方案与 Demo，已把生产化差距列成路线图，不虚构真实上线和真实银行数据。"],
    ], [3300, 6060], font_size=9.7)
    add_callout(doc, "建议收束句", "我刻意没有把金融 Agent 做成只追求回答流畅的聊天框，而是把事实、证据、风险和决策权限拆开。这样系统即使不能做最终审批，也能成为可复核、可迭代的分析基础设施。", fill=GREEN, label_color=GREEN_TEXT)


def chapter_roadmap(doc):
    add_heading(doc, "14. 生产化路线图与结语", 1)
    add_para(doc, "如果把本作品继续推进为可用产品，建议按“数据可信 → 结果可靠 → 流程闭环 → 平台治理”的顺序迭代，而不是一开始就追求更多 Agent。", after=8)
    add_table(doc, ["阶段", "重点工作", "验收标准"], [
        ["P0 数据可信", "接入授权 CRM、制度文档治理、字段脱敏、来源和时间戳", "每条事实和制度依据都可追溯"],
        ["P1 结果可靠", "结构化输出、引用评测、rerank、Golden Dataset 扩充", "事实一致性、引用准确率和拒答率达标"],
        ["P2 流程闭环", "补件任务、人工复核、审批/工单系统、结果回流", "分析结果能转成业务动作并记录处理结果"],
        ["P3 平台治理", "SSO/RBAC、模型网关、观测、成本、CI/CD、容灾", "可审计、可运维、可扩展、可控成本"],
    ], [1500, 5100, 2760], font_size=9.7)
    add_heading(doc, "结语", 2)
    add_para(doc, "这套系统的价值不在于“让模型替银行做决定”，而在于把金融尽调中最耗时、最容易遗漏、最需要证据链的环节变成一个可编排、可解释、可复核的工作台。对 PTA 项目而言，它同时展示了 AI 产品经理从场景判断、架构设计、Prompt/Skill、RAG、MCP 到评测和交付边界的完整能力。", after=9)
    add_callout(doc, "项目定位最终版", "某股份制银行对公信贷尽调智能助手（多智能体系统）——基于公开信息和模拟推演完成的 PTA 方案设计与可运行作品。", fill=LIGHT_BLUE, label_color=DARK_BLUE)


def appendices(doc):
    add_heading(doc, "附录 A：术语表", 1)
    add_table(doc, ["术语", "本项目中的含义"], [
        ["Agent", "在特定职责、工具和规则下完成一类任务的模型驱动节点。"],
        ["Supervisor", "任务主管，负责分类和 Skill 选择，不做最终业务决策。"],
        ["LangGraph", "用于把节点、状态和边组织成可执行工作流的编排框架。"],
        ["MCP", "面向模型的工具协议；本项目通过 /mcp 提供 JSON-RPC 工具演示入口。"],
        ["ToolGateway", "本项目的工具访问边界，负责白名单、Schema、调用和审计。"],
        ["Skill", "版本化的业务规则与执行规范，以 SKILL.md 形式加载并注入上下文。"],
        ["Embedding", "把文本转换为向量表示，用于语义相似度检索。"],
        ["Vector RAG", "切片、向量化、向量库持久化、相似度检索和引用上下文的完整链路。"],
        ["Golden Dataset", "带有期望行为的回归测试样本，覆盖正常、边界和红线场景。"],
        ["人工复核", "系统明确标记结果仅供辅助，需要有权人员继续判断或处理。"],
    ], [2200, 7160], font_size=10)
    add_heading(doc, "附录 B：作品交付清单", 1)
    add_table(doc, ["交付物", "用途"], [
        ["金融多智能体系统_项目说明与系统架构说明书.docx", "本文件：面试答辩与系统整体说明"],
        ["金融多智能体系统_详细操作手册.docx", "按步骤启动、测试和演示系统"],
        ["操作手册.md", "可维护的 Markdown 操作手册源文档"],
        ["README.md", "项目安装、模型、RAG、MCP 和分享说明"],
        ["PORTFOLIO.md", "作品定位、3 分钟演示脚本和生产化差距"],
        ["src/ + web/", "LangGraph、Agent、MCP、RAG、Skill 和前端实现"],
        ["evals/ + tests/", "Golden Dataset、评测脚本和自动化测试"],
    ], [4300, 5060], font_size=10)
    add_callout(doc, "交付提醒", "发给同学或面试官时，建议同时提供项目目录和这两份 Word 文档；不要把 API Key、.venv、Ollama 模型文件或真实敏感数据打包进去。", fill=ORANGE, label_color=ORANGE_TEXT)


def build():
    doc = Document()
    style_document(doc)
    cover(doc)
    add_toc_like(doc)
    chapter_project(doc)
    chapter_positioning(doc)
    chapter_architecture(doc)
    chapter_langgraph(doc)
    chapter_agents(doc)
    chapter_mcp(doc)
    chapter_skill(doc)
    chapter_rag(doc)
    chapter_safety(doc)
    chapter_ui(doc)
    chapter_eval(doc)
    chapter_deploy(doc)
    chapter_interview(doc)
    chapter_roadmap(doc)
    appendices(doc)
    doc.core_properties.title = "某股份制银行对公信贷尽调智能助手：项目说明与系统架构说明书"
    doc.core_properties.subject = "PTA 金融 AI 产品经理项目作品集说明"
    doc.core_properties.author = "AI产品负责人（PTA项目）"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
